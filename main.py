import docx
import requests
from bs4 import BeautifulSoup

def generate_test_paper(questions, testname, file_path):
    doc = docx.Document()
    
    p = doc.add_paragraph()
    runner = p.add_run(testname)
    runner.bold = True

    for i in range(len(questions)):
        question = questions[i]
        k = doc.add_paragraph()
        l = k.add_run("Question"+" "+str(i+1)+":")
        l.bold = True
        question = question.replace("\sqrt{","√")
        question = question.replace("},",",")
        question = question.replace("\[","")
        question = question.replace("\]","")


        doc.add_paragraph(question)
        
    doc.save(file_path)

def scrape_questions():
    response = requests.get("https://www.shaalaa.com/{//url}")
    soup = BeautifulSoup(response.content, 'html.parser')
    questions = soup.find_all(class_="html_wrap")
    questions_text = [question.get_text() for question in questions]

    return questions_text

questionsl = scrape_questions()

testname = "Sample Test"
generate_test_paper(questionsl, testname, "test_paper.docx")
